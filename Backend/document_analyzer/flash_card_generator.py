from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def export_flashcards(cards, docx_filename):
    """
    Creates a DOCX file with flashcards arranged in a two-column table.
    Each row contains a question/term and its corresponding answer/definition.
    
    Args:
        cards (list): List of [front, back] pairs for each flashcard
        docx_filename (str): Output DOCX filename
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Flashcards', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Create table
    table = doc.add_table(rows=len(cards), cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(4)
    
    # Add content and styling to each row
    for i, (front, back) in enumerate(cards):
        # Front of card
        front_cell = table.rows[i].cells[0]
        front_para = front_cell.paragraphs[0]
        front_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        front_run = front_para.add_run(front)
        front_run.font.size = Pt(11)
        front_run.font.bold = True
        
        # Add red background to front cell
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FF9999"/>'.format(nsdecls('w')))
        front_cell._tc.get_or_add_tcPr().append(shading_elm_1)
        
        # Back of card
        back_cell = table.rows[i].cells[1]
        back_para = back_cell.paragraphs[0]
        back_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        back_run = back_para.add_run(back)
        back_run.font.size = Pt(11)
        
        # Add light blue background to back cell
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="CCEBFF"/>'.format(nsdecls('w')))
        back_cell._tc.get_or_add_tcPr().append(shading_elm_2)
        
        # Set row height
        table.rows[i].height = Inches(1.2)
    
    # Add some spacing after the table
    doc.add_paragraph()
    
    # Save the document
    doc.save(docx_filename)