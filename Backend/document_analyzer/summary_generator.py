from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def export_summary(md_string, output_file):
    """
    Exports a markdown-formatted summary to a DOCX file without requiring Pandoc.
    Handles basic markdown syntax like headers, bold text, and bullet points.
    
    Args:
        md_string (str): The markdown-formatted summary text
        output_file (str): Output DOCX filename
    """
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    
    # Split content into lines
    lines = md_string.split('\n')
    
    current_list = False
    
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            current_list = False
            continue
            
        # Handle headers (lines starting with #)
        if line.startswith('#'):
            level = len(line.split()[0])  # Count the number of #
            text = line.lstrip('#').strip()
            heading = doc.add_heading(text, level=min(level, 9))
            current_list = False
            
        # Handle bullet points
        elif line.strip().startswith('*') or line.strip().startswith('-'):
            text = line.lstrip('*- ').strip()
            if not current_list:
                current_list = True
            p = doc.add_paragraph(text)
            p.style = 'List Bullet'
            
        # Handle bold text (text between ** or __)
        else:
            p = doc.add_paragraph()
            current_list = False
            
            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*|__.*?__)', line)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('__') and part.endswith('__'):
                    # Also bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Normal text
                    if part.strip():
                        p.add_run(part)
    
    doc.save(output_file)